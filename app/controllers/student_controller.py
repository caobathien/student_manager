import os
from flask import Blueprint, render_template, session, url_for, flash, redirect, request
from app import db
from app.models.student import Student
from app.models.class_model import Class
from app.models.grade import Grade
from app.models.user import User
from app.models.subject import Subject
from app.models.student_subject import StudentSubject
from app.forms import StudentForm, SearchStudentForm, SubjectRegistrationForm
from flask_login import login_required, current_user
from app.decorators import admin_required, student_required
from sqlalchemy import or_ 
from flask import make_response
import io
from datetime import datetime
import pandas as pd 
from docx import Document 
from docx.shared import Inches 
from io import BytesIO 
from werkzeug.utils import secure_filename

student_bp = Blueprint('student', __name__)

@student_bp.route("/students", methods=['GET', 'POST']) # <-- Thêm POST cho form tìm kiếm
@login_required
def list_students():
    """Hiển thị danh sách sinh viên, có tìm kiếm và lọc."""
    search_form = SearchStudentForm(request.form if request.method == 'POST' else request.args)
    
    # Lấy danh sách lớp để đưa vào dropdown lọc
    classes = Class.query.order_by(Class.name).all()
    # Tạo choices: [(class_id, class_name), ...]
    search_form.class_filter.choices = [('', 'Tất cả các lớp')] + [(str(c.id), c.name) for c in classes]

    # Bắt đầu query, join với Class để có thể sắp xếp/lọc theo tên lớp
    query = Student.query.join(Class).order_by(Class.name, Student.full_name)

    search_term = None
    class_id_filter = None

    # Xử lý tìm kiếm và lọc khi submit form hoặc từ URL
    if request.method == 'POST' and 'submit_search' in request.form and search_form.validate():
         search_term = search_form.search_term.data
         class_id_filter = search_form.class_filter.data
    elif request.method == 'GET': 
        search_term = request.args.get('search_term', None)
        class_id_filter = request.args.get('class_filter', None)
        # Cập nhật giá trị mặc định cho form từ URL
        if search_term:
             search_form.search_term.data = search_term
        if class_id_filter:
             search_form.class_filter.data = class_id_filter

    # Áp dụng bộ lọc tìm kiếm (nếu có)
    if search_term:
        search_pattern = f"%{search_term}%"
        query = query.filter(or_(
            Student.full_name.ilike(search_pattern),
            Student.student_code.ilike(search_pattern)
        ))

    # Áp dụng bộ lọc lớp (nếu có)
    if class_id_filter:
        try:
            query = query.filter(Student.class_id == int(class_id_filter))
        except ValueError:
            pass # Bỏ qua nếu class_id không phải số nguyên

    students = query.all()

    return render_template('students.html',
                           students=students,
                           title="Danh sách sinh viên",
                           search_form=search_form) # <-- Truyền form tìm kiếm ra template


@student_bp.route("/student/add", methods=['GET', 'POST'])
@login_required
@admin_required
def add_student():
    """Thêm một sinh viên mới."""
    form = StudentForm()
    # Populate choices for class_obj
    classes = Class.query.order_by(Class.name).all()
    form.class_obj.choices = [(c.id, c.name) for c in classes]

    # Kiểm tra xem có lớp nào chưa, nếu chưa thì không cho thêm SV
    if not classes:
         flash('Chưa có lớp học nào được tạo. Vui lòng tạo lớp trước khi thêm sinh viên.', 'warning')
         return redirect(url_for('class_admin.list_classes')) # Chuyển đến trang quản lý lớp

    if form.validate_on_submit():
        student = Student(
            student_code=form.student_code.data,
            full_name=form.full_name.data,
            date_of_birth=form.date_of_birth.data,
            gender=form.gender.data,
            # Lấy class_id từ form
            class_id=form.class_obj.data,
            gpa=0.0  # GPA sẽ được tính từ điểm các môn
        )
        db.session.add(student)
        db.session.commit()
        flash('Đã thêm sinh viên mới thành công!', 'success')
        return redirect(url_for('student.list_students'))
    return render_template('student_form.html', title='Thêm sinh viên', form=form, legend='Thêm sinh viên')

@student_bp.route("/student/<int:student_id>/update", methods=['GET', 'POST'])
@login_required
@admin_required
def update_student(student_id):
    """Cập nhật thông tin sinh viên."""
    student = Student.query.get_or_404(student_id)
    form = StudentForm(obj=student)

    # Populate choices for class_obj
    classes = Class.query.order_by(Class.name).all()
    form.class_obj.choices = [(c.id, c.name) for c in classes]

    # Gán sẵn lớp hiện tại cho dropdown khi GET request
    if request.method == 'GET':
        form.class_obj.data = str(student.class_id)

    if form.validate_on_submit():
        student.student_code = form.student_code.data
        student.full_name = form.full_name.data
        student.date_of_birth = form.date_of_birth.data
        student.gender = form.gender.data
        # Lấy class_id từ form
        student.class_id = form.class_obj.data
        # GPA sẽ được tính từ điểm các môn
        db.session.commit()
        flash('Thông tin sinh viên đã được cập nhật!', 'success')
        return redirect(url_for('student.list_students'))

    return render_template('student_form.html', title='Cập nhật sinh viên', form=form, legend='Cập nhật sinh viên')

@student_bp.route("/student/<int:student_id>/delete", methods=['POST'])
@login_required
@admin_required
def delete_student(student_id):
    """Xóa sinh viên."""
    student = Student.query.get_or_404(student_id)
    # Xóa tài khoản người dùng liên kết nếu có
    user = User.query.filter_by(student_id=student.id).first()
    if user:
        db.session.delete(user)
    db.session.delete(student)
    db.session.commit()
    flash('Sinh viên và tài khoản liên kết đã được xóa!', 'success')
    return redirect(url_for('student.list_students'))

@student_bp.route("/students/export/<file_type>") # Route mới nhận file_type
@login_required
def export_students(file_type):
    """Xuất danh sách sinh viên ra file Excel (.xlsx) hoặc Word (.docx)."""

    # Lấy lại logic lọc/tìm kiếm từ URL parameters
    search_term = request.args.get('search_term', None)
    class_id_filter = request.args.get('class_filter', None)

    query = Student.query.join(Class).order_by(Class.name, Student.full_name)

    if search_term:
        search_pattern = f"%{search_term}%"
        query = query.filter(or_(
            Student.full_name.ilike(search_pattern), 
            Student.student_code.ilike(search_pattern)
        ))

    if class_id_filter:
        try:
            query = query.filter(Student.class_id == int(class_id_filter))
        except (ValueError, TypeError): # Bắt cả TypeError nếu class_id_filter là None
            pass 

    students = query.all()

    if not students:
        flash('Không có sinh viên nào để xuất file.', 'info')
        return redirect(url_for('student.list_students', search_term=search_term, class_filter=class_id_filter))

    # --- Tạo tên file ---
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = f"danh_sach_sinh_vien_{timestamp}"

    # --- Xử lý EXCEL (.xlsx) ---
    if file_type == 'xlsx':
        # Chuyển đổi danh sách student objects thành list of dictionaries
        student_data = []
        for student in students:
            student_data.append({
                'Mã SV': student.student_code,
                'Họ và tên': student.full_name,
                'Ngày sinh': student.date_of_birth.strftime('%d-%m-%Y'),
                'Giới tính': student.gender,
                'Lớp': student.class_rel.name if student.class_rel else '',
                'GPA': student.gpa
            })

        # Tạo DataFrame từ list of dictionaries
        df = pd.DataFrame(student_data)

        # Tạo file Excel trong bộ nhớ (BytesIO)
        output = BytesIO()
        writer = pd.ExcelWriter(output, engine='openpyxl')
        df.to_excel(writer, index=False, sheet_name='DanhSachSinhVien')
        # writer.save() # Không cần save với pandas mới
        writer.close() # Thay bằng close()
        output.seek(0) # Đưa con trỏ về đầu stream

        filename = f"{base_filename}.xlsx"
        mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'

    # --- Xử lý WORD (.docx) ---
    elif file_type == 'docx':
        document = Document()
        document.add_heading('Danh sách Sinh viên', 0)

        # Thêm bảng
        table = document.add_table(rows=1, cols=6) # 1 hàng tiêu đề, 6 cột
        table.style = 'Table Grid' # Kiểu bảng có kẻ ô
        table.autofit = False # Tắt tự động vừa cột

        # Set độ rộng cột (tùy chỉnh theo ý bạn)
        widths = (Inches(1.0), Inches(2.0), Inches(1.0), Inches(0.8), Inches(1.0), Inches(0.7))
        hdr_cells = table.rows[0].cells
        header_texts = ['Mã SV', 'Họ và tên', 'Ngày sinh', 'Giới tính', 'Lớp', 'GPA']
        for i, text in enumerate(header_texts):
             hdr_cells[i].text = text
             hdr_cells[i].width = widths[i]

        # Thêm dữ liệu sinh viên
        for student in students:
            row_cells = table.add_row().cells
            row_cells[0].text = student.student_code
            row_cells[1].text = student.full_name
            row_cells[2].text = student.date_of_birth.strftime('%d-%m-%Y')
            row_cells[3].text = student.gender
            row_cells[4].text = student.class_rel.name if student.class_rel else ''
            row_cells[5].text = str(student.gpa)
            # Căn chỉnh độ rộng cho các ô dữ liệu (tùy chọn)
            for i, width in enumerate(widths):
                row_cells[i].width = width

        # Tạo file Word trong bộ nhớ (BytesIO)
        output = BytesIO()
        document.save(output)
        output.seek(0)

        filename = f"{base_filename}.docx"
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    # --- Định dạng không hợp lệ ---
    else:
        flash('Định dạng file không được hỗ trợ.', 'danger')
        return redirect(url_for('student.list_students', search_term=search_term, class_filter=class_id_filter))

    # --- Tạo Response để tải file ---
    response = make_response(output.getvalue())
    response.headers["Content-Disposition"] = f"attachment; filename=\"{filename}\"" # Thêm dấu ngoặc kép cho tên file
    response.headers["Content-type"] = mimetype
    return response

# lấy dữ liệu từ file
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'xlsx', 'csv'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@student_bp.route('/students/import', methods=['GET', 'POST'])
@login_required
@admin_required
def import_students():
    from app.forms import ImportStudentForm
    form = ImportStudentForm()
    if request.method == 'POST':

        # === Nút XEM TRƯỚC ===
        if 'preview' in request.form:
            file = request.files.get('file')
            if not file or file.filename == '':
                flash('⚠️ Vui lòng chọn file!', 'warning')
                return render_template('import_students.html', form=form)

            if allowed_file(file.filename):
                filename = secure_filename(file.filename)
                os.makedirs(UPLOAD_FOLDER, exist_ok=True)
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                file.save(filepath)

                try:
                    if filename.endswith('.xlsx'):
                        df = pd.read_excel(filepath)
                    else:
                        df = pd.read_csv(filepath)

                    # Lưu đường dẫn file vào session để dùng khi confirm
                    session['import_file'] = filepath

                    # Xem trước 5 dòng đầu tiên
                    preview_data = df.head(5).to_html(classes="table table-bordered table-striped table-sm", index=False)
                    return render_template('import_students.html', preview=preview_data, form=form)

                except Exception as e:
                    flash(f'❌ Lỗi khi đọc file: {e}', 'danger')
                    return render_template('import_students.html', form=form)
            else:
                flash('❌ Chỉ hỗ trợ file .xlsx hoặc .csv!', 'danger')
                return render_template('import_students.html', form=form)

        # === Nút XÁC NHẬN NHẬP ===
        elif 'confirm' in request.form:
            filepath = session.get('import_file')
            if not filepath or not os.path.exists(filepath):
                flash('❌ Không tìm thấy file để nhập. Vui lòng tải lại!', 'danger')
                return render_template('import_students.html', form=form)

            try:
                # Đọc lại file
                if filepath.endswith('.xlsx'):
                    df = pd.read_excel(filepath)
                else:
                    df = pd.read_csv(filepath)

                # Chuẩn hóa tên cột
                df.columns = df.columns.str.strip().str.lower()

                added_count = 0

                for _, row in df.iterrows():
                    student_code = str(row.get("mã sv", "")).strip()
                    full_name = str(row.get("họ và tên", "")).strip()
                    date_of_birth_str = str(row.get("ngày sinh", "")).strip()
                    gender = str(row.get("giới tính", "")).strip()
                    class_name = str(row.get("lớp", "")).strip()
                    gpa_str = str(row.get("gpa", "0")).strip()

                    # Bỏ qua dòng thiếu dữ liệu chính
                    if not student_code or not full_name:
                        continue

                    # Chuyển đổi ngày sinh
                    try:
                        date_of_birth = datetime.strptime(date_of_birth_str, "%d-%m-%Y").date()
                    except:
                        try:
                            date_of_birth = datetime.strptime(date_of_birth_str, "%d/%m/%Y").date()
                        except:
                            date_of_birth = None

                    # Chuyển GPA sang float
                    try:
                        gpa = float(gpa_str)
                    except:
                        gpa = 0.0

                    # Tìm lớp trong DB
                    class_obj = Class.query.filter_by(name=class_name).first()

                    # Kiểm tra trùng mã SV
                    if Student.query.filter_by(student_code=student_code).first():
                        continue

                    # Tạo sinh viên mới
                    new_student = Student(
                        student_code=student_code,
                        full_name=full_name,
                        date_of_birth=date_of_birth,
                        gender=gender,
                        class_id=class_obj.id if class_obj else None,
                        gpa=gpa
                    )

                    db.session.add(new_student)
                    added_count += 1

                db.session.commit()
                flash(f"✅ Đã nhập thành công {added_count} sinh viên!", "success")

            except Exception as e:
                db.session.rollback()
                flash(f'❌ Lỗi khi xử lý file: {e}', 'danger')
                return render_template('import_students.html', form=form)

    # GET request hoặc render lại
    return render_template('import_students.html', form=form)

@student_bp.route('/student/<int:student_id>/register_subjects', methods=['GET', 'POST'])
@login_required
def register_subjects(student_id):
    """Đăng ký môn học cho sinh viên."""
    student = Student.query.get_or_404(student_id)
    form = SubjectRegistrationForm()

    # Lấy danh sách môn học chưa đăng ký
    registered_subject_ids = [ss.subject_id for ss in StudentSubject.query.filter_by(student_id=student.id).all()]
    subjects = Subject.query.filter(~Subject.id.in_(registered_subject_ids)).all()
    form.subjects.choices = [(subj.id, f"{subj.code} - {subj.name}") for subj in subjects]

    if form.validate_on_submit():
        selected_subject_ids = form.subjects.data
        # Đăng ký môn học cho sinh viên
        for subject_id in selected_subject_ids:
            # Kiểm tra nếu đã đăng ký rồi thì bỏ qua
            existing = StudentSubject.query.filter_by(student_id=student.id, subject_id=subject_id).first()
            if not existing:
                student_subject = StudentSubject(student_id=student.id, subject_id=subject_id)
                db.session.add(student_subject)
        db.session.commit()
        flash(f'Đã đăng ký {len(selected_subject_ids)} môn học cho sinh viên {student.full_name}!', 'success')
        return redirect(url_for('student.list_students'))

    return render_template('register_subjects.html', title='Đăng ký môn học', form=form, student=student)

@student_bp.route('/my_info')
@login_required
@student_required
def view_my_info():
    """Sinh viên xem thông tin cá nhân và điểm số."""
    # Chỉ cho phép xem thông tin của chính mình
    if not current_user.student:
        flash('Bạn không có quyền truy cập trang này.', 'danger')
        return redirect(url_for('main.home'))
    student = current_user.student
    grades = Grade.query.filter_by(student_id=student.id).all()
    return render_template('student_dashboard.html', title='Thông tin cá nhân', student=student, grades=grades)

@student_bp.route('/my_grades')
@login_required
@student_required
def view_my_grades():
    """Sinh viên xem điểm của mình."""
    # Chỉ cho phép xem điểm của chính mình
    if not current_user.student:
        flash('Bạn không có quyền truy cập trang này.', 'danger')
        return redirect(url_for('main.home'))
    student = current_user.student
    # Lấy tất cả môn học đã đăng ký
    registered_subjects = StudentSubject.query.filter_by(student_id=student.id).all()
    # Tạo danh sách môn học với điểm tương ứng
    subject_grades = []
    for reg in registered_subjects:
        grade = Grade.query.filter_by(student_id=student.id, subject_id=reg.subject_id).first()
        subject_grades.append((reg.subject, grade))
    return render_template('student_grades.html', title='Điểm của tôi', subject_grades=subject_grades)
